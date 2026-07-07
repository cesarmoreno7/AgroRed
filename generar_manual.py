"""
Generador del Manual de Usuario - AgroRed
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Paleta de colores AgroRed ────────────────────────────────────────────────
VERDE_OSCURO  = RGBColor(0x1A, 0x5C, 0x2E)   # #1A5C2E
VERDE_MEDIO   = RGBColor(0x2E, 0x7D, 0x32)   # #2E7D32
VERDE_CLARO   = RGBColor(0x81, 0xC7, 0x84)   # #81C784
NARANJA       = RGBColor(0xE6, 0x5C, 0x00)   # #E65C00
GRIS_OSCURO   = RGBColor(0x21, 0x21, 0x21)   # #212121
GRIS_CLARO    = RGBColor(0xF5, 0xF5, 0xF5)   # #F5F5F5
BLANCO        = RGBColor(0xFF, 0xFF, 0xFF)
AZUL_INFO     = RGBColor(0x01, 0x57, 0x9B)   # #01579B


def set_cell_bg(cell, hex_color: str):
    """Pinta el fondo de una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Línea separadora verde."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A5C2E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = VERDE_OSCURO
    run.font.name = 'Calibri'
    add_horizontal_rule(doc)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = VERDE_MEDIO
    run.font.name = 'Calibri'
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NARANJA
    run.font.name = 'Calibri'
    return p


def body_text(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_OSCURO
    run.font.name = 'Calibri'
    run.italic = italic
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_OSCURO
    run.font.name = 'Calibri'
    return p


def note_box(doc, text):
    """Caja de nota con fondo gris claro."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), '81C784')
        pBdr.append(el)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F8F2')
    pPr.append(shd)
    run = p.add_run('📌  ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = VERDE_OSCURO
    run.font.name = 'Calibri'
    return p


# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# ─── Márgenes ─────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AGRO红 PLATAFORMA DIGITAL')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = VERDE_OSCURO
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema Territorial de Seguridad Alimentaria')
run.font.size = Pt(16)
run.font.color.rgb = NARANJA
run.font.name = 'Calibri'
run.italic = True

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL DE USUARIO')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = VERDE_MEDIO
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in [
    ('Versión:', '1.0'),
    ('Fecha:', 'Mayo 2026'),
    ('Clasificación:', 'Uso Interno'),
    ('Proyecto:', 'AgroRed – GovTech / FoodTech Colombia'),
]:
    run = p.add_run(f'{label}  ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_OSCURO
    run.font.name = 'Calibri'
    run2 = p.add_run(value + '\n')
    run2.font.size = Pt(11)
    run2.font.color.rgb = GRIS_OSCURO
    run2.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '1.  Introducción')
body_text(doc,
    'AgroRed es una plataforma digital de gestión territorial orientada a la seguridad '
    'alimentaria en Colombia. Conecta a productores rurales, operadores logísticos, '
    'municipios, cocinas comunitarias, supermercados y analistas territoriales en un '
    'ecosistema único con más de 15 módulos funcionales integrados.')

body_text(doc,
    'El presente Manual de Usuario describe los perfiles de acceso al sistema, los módulos '
    'disponibles y las actividades que puede realizar cada tipo de usuario. Está dirigido a '
    'funcionarios municipales, operadores de campo y personal técnico que interactúa '
    'diariamente con la plataforma.')

note_box(doc,
    'Para ingresar al sistema, cada usuario requiere credenciales (correo electrónico y '
    'contraseña) asignadas por el administrador municipal de su municipio.')

# ═══════════════════════════════════════════════════════════════════════════════
# 2. DESCRIPCIÓN GENERAL DEL SISTEMA
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '2.  Descripción General del Sistema')
body_text(doc,
    'AgroRed opera como una plataforma consolidada expuesta a través de un API Gateway '
    'central: todos los módulos comparten un mismo backend y una misma base de datos, lo '
    'que garantiza consistencia entre módulos. Cada módulo está diseñado para cubrir un '
    'proceso específico del ciclo alimentario territorial.')

heading2(doc, '2.1  Módulos de la plataforma')
body_text(doc, 'La plataforma se compone de los siguientes módulos:')

MODULES = [
    ('Gestión de Usuarios',      'Autenticación, perfiles y administración de roles.'),
    ('Productores',              'Registro, perfiles y categorización de productores rurales.'),
    ('Ofertas',                  'Publicación de oferta alimentaria disponible para la venta o donación.'),
    ('Demandas',                 'Registro de necesidades alimentarias institucionales (PAE, comedores, hospitales).'),
    ('Rescates Alimentarios',    'Gestión de excedentes agrícolas para redistribución y reducción de desperdicio.'),
    ('Subastas',                 'Subastas agrícolas con 6 algoritmos propietarios (holandesa, inglesa, VICKREY, etc.).'),
    ('Inventario',               'Control de existencias, trazabilidad por lote y alertas de vencimiento.'),
    ('Logística',                'Planificación de rutas óptimas, rastreo GPS y geovallas (VRP).'),
    ('Entregas de Productos',    'Trazabilidad de la entrega física de productos a instituciones: fecha, lugar, receptor y detalle por producto.'),
    ('Instituciones',            'Gestión de instituciones demandantes: comedores, colegios, hospitales y otros programas alimentarios.'),
    ('Incidentes',               'Reporte y seguimiento de incidentes territoriales con SLA y alertas.'),
    ('Analítica y Mapas',        'Índice IRAT, observatorio alimentario, mapas GIS y exportación de reportes.'),
    ('Inteligencia Artificial',  'Copiloto conversacional y recomendaciones heurísticas para soporte a la decisión.'),
    ('Notificaciones',           'Alertas multicanal: correo, SMS, push, in-app y webhook.'),
    ('Automatización',           'Flujos de trabajo automatizados y orquestación de operaciones.'),
    ('Tablas Maestras Territoriales', 'Catálogo de departamentos, municipios, corregimientos y veredas usado como referencia geográfica en el resto de módulos.'),
    ('Mapa de Madurez',          'Panel que resume qué tan completo está cada módulo del ecosistema (datos cargados, endpoints activos).'),
    ('Portal Web / Dashboard',   'Interfaz visual unificada accesible según el rol del usuario.'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
hdr[0].text = 'Módulo'
hdr[1].text = 'Descripción'
for cell in hdr:
    set_cell_bg(cell, '1A5C2E')
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = BLANCO
            r.font.size = Pt(10)
            r.font.name = 'Calibri'

for i, (name, desc) in enumerate(MODULES):
    row = tbl.add_row()
    row.cells[0].text = name
    row.cells[1].text = desc
    bg = 'F1F8F2' if i % 2 == 0 else 'FFFFFF'
    for c in row.cells:
        set_cell_bg(c, bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'

tbl.columns[0].width = Cm(5.5)
tbl.columns[1].width = Cm(10.5)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. PERFILES DE USUARIO
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '3.  Perfiles de Usuario')
body_text(doc,
    'AgroRed define seis perfiles de usuario. Cada perfil representa un actor del ecosistema '
    'alimentario territorial con responsabilidades y accesos diferenciados.')

ROLES = [
    {
        'nombre':   'Administrador Municipal  (admin_municipal)',
        'color':    '1A5C2E',
        'icono':    '🏛️',
        'desc':     (
            'Es el usuario con mayor nivel de acceso dentro de cada municipio. Generalmente '
            'corresponde a un funcionario de la Secretaría de Agricultura o de Desarrollo '
            'Económico municipal. Tiene control total sobre la plataforma en su jurisdicción.'
        ),
        'actividades': [
            'Crear, editar y desactivar cuentas de usuario para todos los demás roles.',
            'Configurar y supervisar todos los módulos de la plataforma.',
            'Publicar y cerrar subastas agrícolas.',
            'Registrar productores y validar su información.',
            'Emitir notificaciones masivas a todos los actores.',
            'Crear y supervisar flujos de automatización.',
            'Acceder a todos los reportes de analítica e IRAT.',
            'Consultar el registro de auditoría del sistema.',
            'Gestionar inventarios y operaciones logísticas.',
            'Reportar y gestionar incidentes territoriales.',
        ],
        'modulos': [
            'Todos los módulos sin restricción.',
        ],
    },
    {
        'nombre':   'Analista Territorial  (territorial_analyst)',
        'color':    '01579B',
        'icono':    '🗺️',
        'desc':     (
            'Profesional técnico (economista, geógrafo, ingeniero agrónomo) encargado del '
            'monitoreo y análisis del territorio alimentario. Su acceso es principalmente '
            'de lectura para garantizar la integridad de los datos de campo.'
        ),
        'actividades': [
            'Consultar el mapa GIS territorial y capas de datos geoespaciales.',
            'Visualizar reportes del Índice de Riesgo Alimentario Territorial (IRAT).',
            'Acceder a resúmenes estadísticos y tendencias de analítica.',
            'Consultar el panel de inteligencia artificial y recomendaciones.',
            'Ver el listado de productores, ofertas, rescates, demandas y subastas (solo lectura).',
            'Consultar el estado del inventario y las rutas logísticas activas.',
            'Reportar y consultar incidentes territoriales.',
            'Recibir y revisar notificaciones del sistema.',
            'Monitorear flujos de automatización (solo lectura).',
        ],
        'modulos': [
            'Analítica y Mapas — Lectura completa.',
            'Inteligencia Artificial — Lectura.',
            'Productores — Lectura.',
            'Ofertas — Lectura.',
            'Rescates — Lectura.',
            'Demandas — Lectura.',
            'Inventario — Lectura.',
            'Logística — Lectura.',
            'Incidentes — Lectura y creación.',
            'Subastas — Lectura.',
            'Notificaciones — Lectura.',
            'Automatización — Lectura.',
        ],
    },
    {
        'nombre':   'Productor  (producer)',
        'color':    '2E7D32',
        'icono':    '🌱',
        'desc':     (
            'Agricultor o asociación de productores rurales. Es el actor primario que '
            'incorpora la oferta alimentaria al sistema. Puede ser individual o representar '
            'a una cooperativa o asociación campesina.'
        ),
        'actividades': [
            'Registrar y actualizar su perfil de productor (cultivos, ubicación, capacidad).',
            'Publicar ofertas de productos agrícolas disponibles.',
            'Registrar excedentes para rescate alimentario.',
            'Participar en subastas publicando lotes de productos.',
            'Consultar el mapa GIS para ver demandas cercanas a su área.',
            'Recibir notificaciones sobre pujas y adjudicaciones.',
        ],
        'modulos': [
            'Productores — Registro y gestión de su propio perfil.',
            'Ofertas — Creación y gestión de sus ofertas.',
            'Rescates — Creación y gestión de sus rescates.',
            'Subastas — Publicación de lotes y seguimiento.',
            'Analítica / Mapas — Solo vista de mapa.',
        ],
    },
    {
        'nombre':   'Cocina Comunitaria  (community_kitchen)',
        'color':    'E65C00',
        'icono':    '🍲',
        'desc':     (
            'Representa comedores comunitarios, ollas populares, restaurantes escolares '
            '(PAE) y hospitales que requieren suministro alimentario regular. Su función '
            'principal es registrar la demanda institucional y adquirir productos.'
        ),
        'actividades': [
            'Registrar demandas alimentarias institucionales (tipo de alimento, cantidad, fecha requerida).',
            'Consultar la oferta agrícola disponible.',
            'Consultar rescates alimentarios disponibles para beneficiarse.',
            'Participar en subastas realizando pujas y aceptando precio holandés.',
            'Ver el mapa GIS para identificar productores y ofertas cercanas.',
            'Recibir notificaciones de adjudicación y logística.',
        ],
        'modulos': [
            'Demandas — Creación y gestión de sus demandas.',
            'Ofertas — Lectura.',
            'Rescates — Lectura y participación.',
            'Subastas — Participación (pujar, aceptar precio).',
            'Analítica / Mapas — Solo vista de mapa.',
        ],
    },
    {
        'nombre':   'Operador Logístico  (logistics_operator)',
        'color':    '4A148C',
        'icono':    '🚚',
        'desc':     (
            'Transportista, empresa de logística o comprador comercial (supermercado mayorista) '
            'que gestiona el traslado físico de los productos. También puede actuar como '
            'comprador en subastas.'
        ),
        'actividades': [
            'Crear y gestionar órdenes de transporte y rutas de distribución.',
            'Actualizar el estado de entregas en tiempo real (GPS / geovallas).',
            'Gestionar el inventario asociado a su operación.',
            'Reportar incidentes durante el transporte.',
            'Participar en subastas realizando pujas.',
            'Consultar demandas activas para planificar operaciones.',
            'Ver ofertas y rescates disponibles para coordinar recogidas.',
        ],
        'modulos': [
            'Logística — Gestión completa.',
            'Inventario — Gestión completa.',
            'Incidentes — Creación y consulta.',
            'Subastas — Participación (pujar, aceptar precio).',
            'Ofertas — Lectura.',
            'Rescates — Lectura.',
            'Demandas — Lectura.',
            'Analítica / Mapas — Solo vista de mapa.',
        ],
    },
    {
        'nombre':   'Supermercado  (supermarket)',
        'color':    '880E4F',
        'icono':    '🛒',
        'desc':     (
            'Cadena de distribución minorista o mayorista que participa como comprador '
            'de productos agrícolas y puede registrar oferta de productos. '
            'Acceso orientado a operaciones comerciales.'
        ),
        'actividades': [
            'Registrar ofertas de productos para la plataforma.',
            'Consultar la oferta agrícola disponible.',
            'Acceder a resúmenes de analítica para planificación comercial.',
            'Recibir notificaciones de disponibilidad.',
        ],
        'modulos': [
            'Ofertas — Creación y lectura.',
            'Analítica — Lectura de resúmenes.',
        ],
    },
    {
        'nombre':   'Agente de Monitoreo  (monitoring_agent)',
        'color':    '546E7A',
        'icono':    '🛰️',
        'desc':     (
            'Personal de campo o de control territorial encargado de vigilar el estado '
            'operativo de la red y reportar incidencias. A diferencia de los demás roles, '
            'no tiene un módulo exclusivo asignado: su acceso corresponde al de cualquier '
            'usuario autenticado sobre los módulos que no tienen una restricción de rol '
            'explícita, con foco en el registro y seguimiento de incidentes.'
        ),
        'actividades': [
            'Reportar incidentes territoriales detectados en campo.',
            'Consultar el estado de incidentes previamente reportados.',
            'Revisar notificaciones operativas del sistema.',
        ],
        'modulos': [
            'Incidentes — Reporte y consulta.',
            'Notificaciones — Recepción.',
        ],
    },
]

for role in ROLES:
    doc.add_paragraph()
    # Encabezado del rol
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(f"  {role['icono']}  {role['nombre']}  ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLANCO
    run.font.name = 'Calibri'
    # Fondo de color
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), role['color'])
    pPr.append(shd)

    body_text(doc, role['desc'])

    heading3(doc, 'Actividades principales')
    for act in role['actividades']:
        bullet(doc, act)

    heading3(doc, 'Módulos accesibles')
    for mod in role['modulos']:
        bullet(doc, mod)

    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. MATRIZ DE ACCESO POR MÓDULO
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '4.  Menú del Sistema por Perfil de Usuario')
body_text(doc,
    'El menú lateral del portal web no es igual para todos los usuarios: cada perfil ve '
    'únicamente las secciones y opciones para las que tiene acceso. Esta sección muestra, '
    'agrupado exactamente como aparece en el menú lateral, qué opciones ve cada perfil al '
    'iniciar sesión.')
note_box(doc,
    'El Tablero Institucional, el Mapa de Madurez y las cuatro páginas de Tablas Maestras '
    '(Departamentos, Municipios, Corregimientos, Veredas) son visibles para cualquier usuario '
    'autenticado, sin importar su perfil.')

# Fuente: apps/web-dashboard/src/components/Layout.tsx (MENU_GROUPS) +
# apps/web-dashboard/src/config/moduleAccess.ts (MODULE_ACCESS). Si alguno de esos dos
# archivos cambia, esta sección debe actualizarse para que el manual siga reflejando el
# menú real.
MENU_ROLES_COLS   = ['Admin\nMunicipal', 'Analista\nTerritorial', 'Productor', 'Cocina\nComunitaria', 'Operador\nLogístico', 'Super-\nmercado', 'Agente de\nMonitoreo*']
MENU_ROLE_COLORS  = ['1A5C2E', '01579B', '2E7D32', 'E65C00', '4A148C', '880E4F', '546E7A']
MENU_ROLE_KEYS    = ['admin_municipal', 'territorial_analyst', 'producer', 'community_kitchen', 'logistics_operator', 'supermarket', 'monitoring_agent']

MODULE_ACCESS = {
    'admin_municipal': {'user-service','producer-service','offer-service','rescue-service','demand-service',
                         'institution-service','origins-service','inventory-service','logistics-service',
                         'incident-service','notification-service','auction-service','analytics-service',
                         'ml-service','catalog-service','delivery-service'},
    'territorial_analyst': {'producer-service','offer-service','rescue-service','demand-service',
                             'institution-service','logistics-service','incident-service','auction-service',
                             'analytics-service','ml-service','catalog-service','delivery-service'},
    'logistics_operator': {'producer-service','offer-service','rescue-service','demand-service',
                            'inventory-service','logistics-service','incident-service','notification-service',
                            'auction-service','catalog-service','delivery-service'},
    'community_kitchen': {'offer-service','rescue-service','demand-service','auction-service','delivery-service'},
    'producer': {'producer-service','rescue-service','offer-service','auction-service','incident-service','delivery-service'},
    'supermarket': {'offer-service','auction-service'},
    'monitoring_agent': set(),  # sin entrada en moduleAccess.ts — ver nota más abajo
}

MENU_SECTIONS = [
    ('Principal', [('Tablero Institucional', None)]),
    ('Territorio', [
        ('Mapa Territorial',   'logistics-service'),
        ('Instituciones',      'institution-service'),
        ('Orígenes Aliados',   'origins-service'),
    ]),
    ('Operaciones', [
        ('Productores',              'producer-service'),
        ('Ofertas',                  'offer-service'),
        ('Rescates',                 'rescue-service'),
        ('Demandas',                 'demand-service'),
        ('Inventario',               'inventory-service'),
        ('Flota en tiempo real',     'logistics-service'),
        ('Geocercas logísticas',     'logistics-service'),
        ('Entregas de Productos',    'delivery-service'),
    ]),
    ('Administración', [
        ('Usuarios',        'user-service'),
        ('Canales',         'catalog-service'),
        ('Organizaciones',  'catalog-service'),
        ('Productos',       'catalog-service'),
        ('Categorías',      'catalog-service'),
    ]),
    ('Tablas Maestras', [
        ('Departamentos',    None),
        ('Municipios',       None),
        ('Corregimientos',   None),
        ('Veredas',          None),
    ]),
    ('Inteligencia', [
        ('Incidencias Sociales',  'incident-service'),
        ('Notificaciones',        'notification-service'),
        ('Subastas',              'auction-service'),
        ('Alertas IRAT',          'analytics-service'),
        ('Apoyo a Decisión',      'ml-service'),
        ('Copiloto IA',           'user-service'),
        ('Mapa de Madurez',       None),
    ]),
]

for section_label, items in MENU_SECTIONS:
    heading2(doc, section_label)
    cols = 1 + len(MENU_ROLES_COLS)
    tbl_sec = doc.add_table(rows=1, cols=cols)
    tbl_sec.style = 'Table Grid'
    tbl_sec.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl_sec.rows[0].cells
    hdr[0].text = 'Opción de menú'
    set_cell_bg(hdr[0], '1A5C2E')
    for r in hdr[0].paragraphs[0].runs:
        r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(8.5); r.font.name = 'Calibri'
    for i, (label, color) in enumerate(zip(MENU_ROLES_COLS, MENU_ROLE_COLORS)):
        hdr[i + 1].text = label
        set_cell_bg(hdr[i + 1], color)
        for r in hdr[i + 1].paragraphs[0].runs:
            r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(8); r.font.name = 'Calibri'

    for j, (item_label, module) in enumerate(items):
        row = tbl_sec.add_row()
        row.cells[0].text = item_label
        bg = 'F1F8F2' if j % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        for r in row.cells[0].paragraphs[0].runs:
            r.font.size = Pt(9.5); r.font.name = 'Calibri'; r.bold = True

        for k, role_key in enumerate(MENU_ROLE_KEYS, start=1):
            visible = module is None or module in MODULE_ACCESS[role_key]
            row.cells[k].text = '✔' if visible else '✘'
            set_cell_bg(row.cells[k], bg)
            for r in row.cells[k].paragraphs[0].runs:
                r.font.size = Pt(10.5); r.font.name = 'Calibri'
                r.font.color.rgb = VERDE_MEDIO if visible else RGBColor(0xBD, 0xBD, 0xBD)
            row.cells[k].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    tbl_sec.columns[0].width = Cm(4.2)
    for i in range(1, cols):
        tbl_sec.columns[i].width = Cm(1.85)
    doc.add_paragraph()

note_box(doc,
    '* Agente de Monitoreo: en la versión actual de la plataforma este perfil no tiene módulos '
    'configurados en el menú — solo ve las opciones que son visibles para cualquier usuario '
    'autenticado (Tablero, Mapa de Madurez y Tablas Maestras). Se recomienda validar con el '
    'equipo técnico si esto es el comportamiento previsto o si falta completar su configuración '
    'de acceso.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. DESCRIPCIÓN DETALLADA DE MÓDULOS
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '5.  Descripción Detallada de Módulos')

MODULES_DETAIL = [
    (
        '5.1  Gestión de Usuarios',
        ('Permite al Administrador Municipal crear, editar y desactivar cuentas de usuario. '
         'Desde aquí se asignan los roles y se vinculan los usuarios a un municipio (tenant). '
         'Incluye el registro de nuevos usuarios y la gestión de contraseñas.'),
        ['Crear nuevas cuentas de usuario.', 'Asignar o cambiar el rol de un usuario.',
         'Desactivar cuentas sin eliminar el historial.', 'Consultar el listado de usuarios activos.',
         'Restablecer contraseñas.'],
        'Exclusivo del Administrador Municipal.'
    ),
    (
        '5.2  Productores',
        ('Repositorio de todos los productores rurales registrados en la plataforma. '
         'Permite gestionar perfiles detallados con información geográfica, tipos de cultivo, '
         'capacidad productiva y datos de contacto. Soporta importación masiva por CSV.'),
        ['Registrar nuevo productor (individuo o asociación).', 'Editar datos del perfil.',
         'Importar productores en lote mediante archivo CSV.', 'Categorizar por tipo de producto y zona geográfica.',
         'Consultar el listado con filtros por municipio, producto o estado.'],
        'Administrador Municipal y Productor (su propio perfil). Analista y Logístico pueden consultar.'
    ),
    (
        '5.3  Ofertas',
        ('Módulo central de publicación de oferta alimentaria. Los productores y supermercados '
         'publican productos disponibles con precio, cantidad, fecha de disponibilidad y '
         'ubicación. El sistema realiza matching automático con demandas registradas.'),
        ['Publicar una oferta con descripción, precio, cantidad y fecha.', 'Editar o retirar una oferta activa.',
         'Consultar el listado de ofertas con filtros avanzados.', 'Ver el matching automático con demandas.',
         'Notificar a interesados cuando una nueva oferta es publicada.'],
        'Publicación: Productor, Supermercado, Admin. Consulta: todos los roles.'
    ),
    (
        '5.4  Demandas',
        ('Registro de necesidades alimentarias institucionales. Las cocinas comunitarias, '
         'comedores escolares y hospitales declaran qué productos necesitan, en qué cantidades '
         'y para cuándo. El sistema busca ofertas que satisfagan la demanda.'),
        ['Registrar una demanda especificando producto, cantidad, fecha y ubicación de entrega.',
         'Actualizar el estado de la demanda (pendiente, en proceso, satisfecha).',
         'Consultar el historial de demandas propias.', 'Ver matching automático con ofertas disponibles.'],
        'Creación: Cocina Comunitaria y Admin. Consulta: Admin, Analista, Logístico.'
    ),
    (
        '5.5  Rescates Alimentarios',
        ('Gestión de excedentes agrícolas que de otro modo serían desperdiciados. Los '
         'productores registran sus excedentes y el sistema conecta con cocinas comunitarias '
         'y organizaciones que puedan aprovecharlos, priorizando zonas de vulnerabilidad.'),
        ['Registrar un rescate alimentario indicando tipo, cantidad y fecha de disponibilidad.',
         'Asignar un rescate a una cocina comunitaria o beneficiario.',
         'Consultar el historial de rescates y métricas de toneladas rescatadas.',
         'Generar reporte de impacto ambiental y social.'],
        'Creación: Productor, Cocina Comunitaria, Admin. Consulta: Analista, Logístico.'
    ),
    (
        '5.6  Subastas',
        ('Sistema de subastas agrícolas con 6 algoritmos propietarios: inglesa (puja '
         'ascendente), holandesa (precio descendente), VICKREY (sobre sellado), '
         'combinatoria, inversa y de reserva. Permite transacciones transparentes '
         'entre productores y compradores institucionales o logísticos.'),
        ['Publicar un lote de subasta con precio inicial, algoritmo y fecha de cierre.',
         'Realizar una puja sobre un lote activo.', 'Aceptar el precio holandés en curso.',
         'Cerrar una subasta manualmente (Admin).', 'Consultar el historial de subastas y adjudicaciones.',
         'Recibir notificaciones de puja superada o adjudicación.'],
        'Publicar: Productor y Admin. Pujar/Aceptar: Cocina Comunitaria, Logístico, Admin. Consultar: todos.'
    ),
    (
        '5.7  Inventario',
        ('Control de existencias por lote con trazabilidad completa desde el origen. '
         'Genera alertas automáticas de productos próximos a vencer y permite '
         'registrar movimientos de entrada y salida.'),
        ['Registrar entrada de productos al inventario con datos de lote y fecha.',
         'Registrar salida o consumo de productos.', 'Consultar existencias actuales por categoría.',
         'Ver alertas de vencimiento próximo (configurable: 3, 7, 15 días).',
         'Generar reporte de rotación de inventario.'],
        'Gestión completa: Logístico y Admin. Consulta: Analista Territorial.'
    ),
    (
        '5.8  Logística',
        ('Módulo de planificación y seguimiento de la cadena de transporte. Utiliza '
         'algoritmos VRP (Vehicle Routing Problem) para optimizar rutas multiparada. '
         'Incluye rastreo GPS en tiempo real y configuración de geovallas de alerta.'),
        ['Crear una orden de transporte asignando origen, destino y producto.',
         'Optimizar automáticamente la ruta para múltiples paradas.',
         'Actualizar el estado de una entrega (en tránsito, entregado, fallido).',
         'Configurar geovallas para alertas de desvío de ruta.',
         'Consultar el historial de entregas y métricas de cumplimiento.'],
        'Gestión completa: Logístico y Admin. Consulta: Analista Territorial.'
    ),
    (
        '5.9  Incidentes',
        ('Registro y seguimiento de incidentes territoriales (plagas, inundaciones, '
         'bloqueos viales, escasez) con sistema de prioridades, SLA de respuesta '
         'y escalamiento automático. Permite coordinar respuestas interinstitucionales.'),
        ['Reportar un incidente indicando tipo, ubicación, gravedad y descripción.',
         'Asignar responsable y fecha límite de resolución (SLA).',
         'Actualizar el estado del incidente (abierto, en gestión, resuelto).',
         'Recibir alertas de incidentes no resueltos dentro del SLA.',
         'Consultar el mapa de incidentes activos en el territorio.'],
        'Creación: Analista Territorial, Logístico, Admin. Consulta: los mismos roles.'
    ),
    (
        '5.10  Analítica y Mapas',
        ('Centro de inteligencia territorial. Calcula el Índice de Riesgo Alimentario '
         'Territorial (IRAT) que combina variables de oferta, demanda, accesibilidad '
         'y vulnerabilidad socioeconómica. Incluye mapas GIS interactivos con capas '
         'configurables y exportación de reportes en Excel/PDF.'),
        ['Visualizar el mapa GIS con capas de productores, demandas, incidentes y zonas de riesgo.',
         'Consultar el IRAT por municipio, vereda o zona geográfica.',
         'Generar reportes periódicos (semanal, mensual, anual) exportables.',
         'Comparar indicadores entre períodos o territorios.',
         'Acceder al observatorio alimentario con tendencias históricas.'],
        'Acceso completo: Admin y Analista Territorial. Solo mapa: Productor, Cocina, Logístico. Resúmenes: Supermercado.'
    ),
    (
        '5.11  Inteligencia Artificial',
        ('Motor de soporte a la decisión basado en heurísticas. Genera recomendaciones '
         'sobre asignación de recursos, predicción de escasez y optimización de '
         'distribución alimentaria.'),
        ['Consultar recomendaciones automáticas para asignación de oferta-demanda.',
         'Revisar clasificaciones de riesgo territorial generadas por el modelo.',
         'Visualizar tendencias predictivas de producción y consumo.'],
        'Acceso: Admin y Analista Territorial (solo lectura).'
    ),
    (
        '5.12  Notificaciones',
        ('Sistema multicanal de alertas configurables: correo electrónico, SMS, '
         'notificaciones push, in-app y webhooks para integración con sistemas externos. '
         'Permite crear plantillas y programar envíos masivos.'),
        ['Crear y enviar notificaciones a uno o varios grupos de usuarios.',
         'Configurar alertas automáticas por eventos del sistema (puja, adjudicación, vencimiento).',
         'Consultar el historial de notificaciones enviadas y su estado de entrega.',
         'Gestionar plantillas de mensajes reutilizables.'],
        'Envío: Admin exclusivo. Lectura: Admin y Analista. Recepción: todos los roles.'
    ),
    (
        '5.13  Automatización',
        ('Orquestador de flujos de trabajo que permite configurar tareas automáticas '
         'basadas en eventos o calendarios: cierre automático de subastas, alertas '
         'de inventario, generación de reportes periódicos, etc.'),
        ['Crear flujos de automatización (triggers + acciones).',
         'Activar, pausar o eliminar una automatización existente.',
         'Consultar el registro de ejecuciones y su resultado.',
         'Recibir alertas cuando una automatización falla.'],
        'Creación y gestión: Admin exclusivo. Consulta: Analista Territorial.'
    ),
    (
        '5.14  Entregas de Productos',
        ('Registra la entrega física de productos a instituciones, con trazabilidad completa: '
         'productor de origen, institución receptora, fecha y lugar de entrega, receptor y '
         'detalle línea por línea de los productos entregados (cantidad, unidad, precio).'),
        ['Registrar una nueva entrega asociando productor, institución y fecha.',
         'Agregar el detalle de productos entregados (uno o varios por entrega).',
         'Consultar el historial de entregas con su número de entrega único (formato ENT-AAAA-####).',
         'Verificar el estado de una entrega (pendiente o recibida).'],
        'Creación y gestión: Admin. Consulta: Analista, Cocina Comunitaria, Logístico.'
    ),
    (
        '5.15  Instituciones',
        ('Directorio de instituciones demandantes: comedores comunitarios, colegios (PAE), '
         'hospitales y otros programas alimentarios. Cada institución queda vinculada a un '
         'municipio y sirve como referencia para registrar demandas y entregas.'),
        ['Registrar una nueva institución con tipo, nombre, contacto y municipio.',
         'Editar los datos de una institución existente.',
         'Cambiar el estado de una institución (activa, inactiva, en revisión).',
         'Consultar el listado de instituciones por municipio o tipo.'],
        'Creación y gestión: Admin exclusivo. Consulta: Analista, Cocina Comunitaria, Logístico.'
    ),
    (
        '5.16  Tablas Maestras Territoriales',
        ('Catálogo de referencia geográfica de Colombia usado por el resto de módulos para '
         'ubicar productores, instituciones, incidentes y entregas: departamentos, municipios, '
         'corregimientos y veredas, con sus códigos DANE.'),
        ['Consultar el listado de departamentos, municipios, corregimientos y veredas.',
         'Registrar o editar una unidad territorial (Admin).',
         'Usar estos catálogos como listas desplegables al registrar productores, instituciones '
         'o incidentes en otros módulos.'],
        'Gestión completa: Admin exclusivo. Consulta (vía listas desplegables): todos los roles.'
    ),
    (
        '5.17  Mapa de Madurez',
        ('Panel de diagnóstico interno que resume qué tan completo está cada módulo del '
         'ecosistema: si tiene datos cargados, si sus endpoints responden correctamente y '
         'cuántos registros existen. Pensado como apoyo para el equipo técnico y el '
         'Administrador Municipal durante la puesta en marcha de un nuevo municipio.'),
        ['Consultar el estado de madurez de cada módulo (usuarios, productores, ofertas, etc.).',
         'Identificar módulos sin datos cargados aún en un municipio nuevo.'],
        'Acceso: Admin Municipal.'
    ),
]

for title, desc, actions, access in MODULES_DETAIL:
    heading2(doc, title)
    body_text(doc, desc)
    heading3(doc, 'Funcionalidades disponibles')
    for a in actions:
        bullet(doc, a)
    note_box(doc, f'Quién puede usar este módulo: {access}')
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. FLUJOS DE TRABAJO POR PERFIL
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '6.  Flujos de Trabajo por Perfil de Usuario')

FLOWS = [
    (
        'Administrador Municipal — Flujo de incorporación de productor',
        [
            'Ingresar al módulo Productores.',
            'Seleccionar "Registrar nuevo productor".',
            'Completar el formulario con datos del productor (nombre, vereda, tipo de cultivo, contacto).',
            'Guardar el registro. El sistema asigna un ID único y notifica al productor.',
            '(Opcional) Crear credenciales de acceso para el productor desde Gestión de Usuarios.',
        ]
    ),
    (
        'Productor — Publicar una oferta de producto',
        [
            'Ingresar al módulo Ofertas.',
            'Seleccionar "Nueva oferta".',
            'Completar: tipo de producto, cantidad disponible, precio unitario, fecha de disponibilidad y ubicación.',
            'Publicar. El sistema realiza matching automático con demandas activas y envía notificaciones.',
            'Monitorear el estado de la oferta (activa, en negociación, cerrada).',
        ]
    ),
    (
        'Cocina Comunitaria — Registrar una demanda alimentaria',
        [
            'Ingresar al módulo Demandas.',
            'Seleccionar "Nueva demanda".',
            'Especificar: tipo de alimento, cantidad requerida, fecha de entrega y punto de recibo.',
            'Guardar. El sistema busca ofertas disponibles y muestra coincidencias.',
            'Seleccionar una oferta coincidente y solicitar vinculación.',
            'Coordinar la entrega con el operador logístico asignado.',
        ]
    ),
    (
        'Operador Logístico — Gestionar una entrega',
        [
            'Ingresar al módulo Logística.',
            'Crear una nueva orden de transporte (origen, destino, producto, fecha).',
            'Usar la optimización automática de ruta (VRP) para múltiples paradas.',
            'Actualizar el estado en tiempo real durante el transporte.',
            'Registrar la entrega como completada y adjuntar evidencia si aplica.',
            'En caso de incidente durante el transporte, reportarlo en el módulo de Incidentes.',
        ]
    ),
    (
        'Analista Territorial — Generar reporte de riesgo alimentario',
        [
            'Ingresar al módulo Analítica y Mapas.',
            'Seleccionar la capa "IRAT" en el mapa GIS.',
            'Configurar los filtros: municipio, período de tiempo y variables a incluir.',
            'Generar el reporte y descargarlo en formato Excel o PDF.',
            'Compartir el informe con la administración municipal o instancias superiores.',
        ]
    ),
    (
        'Administrador Municipal — Registrar una entrega de productos',
        [
            'Ingresar al módulo Entregas de Productos.',
            'Seleccionar "Nueva entrega" y elegir el productor de origen y la institución receptora.',
            'Indicar fecha, lugar de entrega y responsable que recibe.',
            'Agregar el detalle de productos entregados (producto, cantidad, unidad, precio unitario).',
            'Guardar. El sistema genera un número de entrega único (ENT-AAAA-####) para trazabilidad.',
        ]
    ),
]

for title, steps in FLOWS:
    heading2(doc, title)
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after  = Pt(3)
        run1 = p.add_run(f'{i}.  ')
        run1.bold = True
        run1.font.color.rgb = VERDE_MEDIO
        run1.font.size = Pt(10.5)
        run1.font.name = 'Calibri'
        run2 = p.add_run(step)
        run2.font.size = Pt(10.5)
        run2.font.color.rgb = GRIS_OSCURO
        run2.font.name = 'Calibri'
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. AUTENTICACIÓN Y SEGURIDAD
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '7.  Autenticación y Seguridad')
body_text(doc,
    'AgroRed implementa un sistema de autenticación basado en JWT (JSON Web Tokens) con '
    'validez configurable (predeterminado: 8 horas). Al iniciar sesión, el sistema retorna '
    'un token que debe incluirse en cada petición a la plataforma.')

heading2(doc, '7.1  Inicio de sesión')
for step in [
    'Ingresar a la URL de la plataforma.',
    'Proporcionar correo electrónico y contraseña.',
    'El sistema verifica las credenciales y retorna un token de acceso y la lista de módulos habilitados según el rol.',
    'La sesión expira automáticamente al transcurrir el tiempo configurado.',
]:
    bullet(doc, step)

heading2(doc, '7.2  Cierre de sesión')
body_text(doc,
    'Al cerrar sesión, el token es invalidado inmediatamente mediante una lista negra en '
    'caché (Redis). Esto garantiza que tokens interceptados no puedan reutilizarse.')

heading2(doc, '7.3  Multi-tenancy')
body_text(doc,
    'Cada usuario pertenece a un municipio (tenant). Los datos de un municipio no son '
    'visibles para usuarios de otros municipios, garantizando la segregación de información territorial.')

heading2(doc, '7.4  Recomendaciones de seguridad')
for rec in [
    'No compartir credenciales de acceso con otras personas.',
    'Cerrar sesión al terminar de usar la plataforma en equipos compartidos.',
    'Reportar de inmediato al Administrador Municipal cualquier acceso sospechoso.',
    'Usar contraseñas de al menos 8 caracteres con combinación de letras, números y símbolos.',
]:
    bullet(doc, rec)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. GLOSARIO
# ═══════════════════════════════════════════════════════════════════════════════
heading1(doc, '8.  Glosario')

GLOSARIO = [
    ('IRAT', 'Índice de Riesgo Alimentario Territorial. Indicador compuesto que mide el nivel de vulnerabilidad alimentaria de un territorio a partir de variables de oferta, demanda, logística y condiciones socioeconómicas.'),
    ('PAE',  'Programa de Alimentación Escolar. Programa gubernamental colombiano que suministra complementos nutricionales a estudiantes de colegios públicos.'),
    ('VRP',  'Vehicle Routing Problem (Problema de Enrutamiento de Vehículos). Algoritmo de optimización para planificar rutas de distribución minimizando distancia o tiempo.'),
    ('JWT',  'JSON Web Token. Estándar de autenticación que permite verificar la identidad del usuario sin mantener sesiones en el servidor.'),
    ('GIS',  'Geographic Information System. Sistema de información geográfica para visualizar y analizar datos con componente espacial/territorial.'),
    ('SLA',  'Service Level Agreement. Acuerdo de nivel de servicio; en el contexto de incidentes, define el tiempo máximo aceptable para resolver un evento.'),
    ('Tenant', 'En AgroRed, corresponde a un municipio. Cada tenant tiene sus propios datos aislados de otros municipios en la plataforma.'),
    ('Subasta holandesa', 'Modalidad de subasta donde el precio comienza alto y desciende progresivamente hasta que un comprador acepta.'),
    ('Subasta VICKREY', 'Subasta de sobre sellado donde gana el mayor postor pero paga el segundo precio más alto.'),
    ('Matching automático', 'Proceso del sistema que cruza automáticamente ofertas y demandas para identificar pares de interés compatible.'),
    ('Rescate alimentario', 'Proceso de recuperar excedentes de producción que de otro modo serían desperdiciados, para redistribuirlos a población vulnerable.'),
]

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = 'Table Grid'
hdr3 = tbl3.rows[0].cells
hdr3[0].text = 'Término'
hdr3[1].text = 'Definición'
for cell in hdr3:
    set_cell_bg(cell, '1A5C2E')
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(10); r.font.name = 'Calibri'

for i, (term, defn) in enumerate(GLOSARIO):
    row = tbl3.add_row()
    row.cells[0].text = term
    row.cells[1].text = defn
    bg = 'F1F8F2' if i % 2 == 0 else 'FFFFFF'
    for c in row.cells:
        set_cell_bg(c, bg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10); r.font.name = 'Calibri'

tbl3.columns[0].width = Cm(4.5)
tbl3.columns[1].width = Cm(11.5)

# ─── Pie de página ────────────────────────────────────────────────────────────
doc.add_paragraph()
add_horizontal_rule(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AgroRed – Manual de Usuario v1.1  |  Julio 2026  |  Confidencial – Uso Interno')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)
run.font.name = 'Calibri'
run.italic = True

# ─── Guardar ──────────────────────────────────────────────────────────────────
output = r'd:\xampp\htdocs\AgroRed\manual_usuario_AGRORED.docx'
doc.save(output)
print(f'[OK]  Documento generado: {output}')
